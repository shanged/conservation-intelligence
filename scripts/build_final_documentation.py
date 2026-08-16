"""Build the two final project handoff documents from reviewed Markdown sources."""

from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
TECH_MD = DOCS / "implementation_report.md"
MDC_MD = DOCS / "mdc_evaluation_user_guide.md"
TECH_DOCX = DOCS / "Conservation_Document_Intelligence_Technical_Handoff_Guide.docx"
MDC_DOCX = DOCS / "Conservation_Document_Intelligence_MDC_Evaluation_User_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BLACK = "000000"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def font(size: int, *, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def rounded_box(draw, xy, text, fill, outline, title_font, body_font, subtitle=""):
    draw.rounded_rectangle(xy, radius=22, fill=fill, outline=outline, width=4)
    x1, y1, x2, y2 = xy
    title_box = draw.multiline_textbbox((0, 0), text, font=title_font, spacing=7, align="center")
    title_h = title_box[3] - title_box[1]
    body_h = 0
    if subtitle:
        body_box = draw.multiline_textbbox((0, 0), subtitle, font=body_font, spacing=5, align="center")
        body_h = body_box[3] - body_box[1] + 16
    y = y1 + ((y2 - y1) - title_h - body_h) / 2
    draw.multiline_text(((x1 + x2) / 2, y), text, font=title_font, fill="#17365D", anchor="ma", align="center", spacing=7)
    if subtitle:
        draw.multiline_text(((x1 + x2) / 2, y + title_h + 16), subtitle, font=body_font, fill="#475467", anchor="ma", align="center", spacing=5)


def arrow(draw, start, end, color="#5B708B"):
    draw.line([start, end], fill=color, width=6)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 12), (ex - 18 * direction, ey + 12)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 12, ey - 18 * direction), (ex + 12, ey - 18 * direction)]
    draw.polygon(points, fill=color)


def build_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1050), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    tf = font(34, bold=True)
    bf = font(23)
    heading = font(44, bold=True)
    draw.text((900, 45), "Implemented System Architecture", font=heading, fill="#17365D", anchor="ma")
    boxes = {
        "sources": (70, 155, 395, 315),
        "processing": (520, 155, 850, 315),
        "sqlite": (985, 155, 1320, 315),
        "entities": (70, 440, 395, 605),
        "chroma": (520, 440, 850, 605),
        "wiki": (985, 440, 1320, 605),
        "chat": (1395, 440, 1725, 605),
        "app": (730, 760, 1070, 935),
    }
    rounded_box(draw, boxes["sources"], "Public sources", "#F4F8FC", "#8AA9C5", tf, bf, "35 metadata records\nPDF and public web text")
    rounded_box(draw, boxes["processing"], "Offline processing", "#F4F8FC", "#8AA9C5", tf, bf, "Extraction, page markers,\n800-word chunks")
    rounded_box(draw, boxes["sqlite"], "SQLite evidence", "#EEF5FB", "#5C8DB8", tf, bf, "Documents, chunks, entities,\nrelations and Wiki metadata")
    rounded_box(draw, boxes["entities"], "Structured evidence", "#F7F4FB", "#967BB6", tf, bf, "Rule-based entities,\nrelationships and counts")
    rounded_box(draw, boxes["chroma"], "Local semantic index", "#F2F8F5", "#6EA087", tf, bf, "MiniLM windows and\npersistent Chroma")
    rounded_box(draw, boxes["wiki"], "Evidence-backed Wiki", "#FFF8E8", "#C89E3D", tf, bf, "15 generated Markdown pages")
    rounded_box(draw, boxes["chat"], "Question answering", "#FFF2F1", "#C77B70", tf, bf, "Deterministic fallback +\noptional bounded synthesis")
    rounded_box(draw, boxes["app"], "Streamlit application", "#EAF1F8", "#2E74B5", tf, bf, "Corpus  |  Search  |  Wiki\nChatbot  |  Evaluation")
    arrow(draw, (395, 235), (520, 235))
    arrow(draw, (850, 235), (985, 235))
    arrow(draw, (1150, 315), (1150, 440))
    arrow(draw, (985, 520), (850, 520))
    arrow(draw, (985, 520), (1320, 520))
    arrow(draw, (1320, 520), (1395, 520))
    arrow(draw, (985, 260), (395, 520))
    arrow(draw, (685, 605), (820, 760))
    arrow(draw, (1150, 605), (980, 760))
    arrow(draw, (1510, 605), (1060, 790))
    draw.text((900, 1000), "Retrieval and final citations remain local; hosted startup loads precomputed artifacts.", font=font(24), fill="#475467", anchor="mm")
    image.save(path)


def build_ui_diagram(path: Path) -> None:
    image = Image.new("RGB", (1700, 700), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((850, 48), "How to explore the prototype", font=font(44, bold=True), fill="#17365D", anchor="ma")
    names = [
        ("Corpus", "See what is included"),
        ("Search", "Find relevant passages"),
        ("Wiki", "Explore organized evidence"),
        ("Chatbot", "Ask a conservation question"),
        ("Evaluation", "Review example results"),
    ]
    colors = ["#EFF6FC", "#EEF8F4", "#FFF7E6", "#FFF1EF", "#F5F1FA"]
    outlines = ["#6E9BC1", "#6FA18A", "#C99B36", "#C57970", "#967CB5"]
    x = 70
    boxes = []
    for index, ((name, subtitle), fill, outline) in enumerate(zip(names, colors, outlines)):
        box = (x, 200, x + 270, 430)
        boxes.append(box)
        rounded_box(draw, box, f"{index + 1}. {name}", fill, outline, font(31, bold=True), font(22), subtitle)
        x += 325
    for left, right in zip(boxes, boxes[1:]):
        arrow(draw, (left[2] + 10, 315), (right[0] - 10, 315))
    draw.rounded_rectangle((250, 520, 1450, 640), radius=18, fill="#F4F6F9", outline="#B9C2CF", width=3)
    draw.text((850, 580), "For important conclusions: open the cited source, read the passage, and apply conservation judgment.", font=font(26, bold=True), fill="#17365D", anchor="mm")
    image.save(path)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None, underline=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def style_document(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run(running_label), size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("Page "), size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def add_cover(doc: Document, kicker: str, title: str, subtitle: str, audience: str):
    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run(kicker.upper()), size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(title), size=27, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    set_run_font(p.add_run(subtitle), size=14, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run(audience), size=11, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    set_run_font(p.add_run("Working private deployment verified | August 2026"), size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Conservation document research prototype"), size=10, color=BLUE, bold=True)
    doc.add_page_break()


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|https?://\S+)")


def add_inline(paragraph, text: str):
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), bold=True)
        elif token.startswith("`"):
            set_run_font(paragraph.add_run(token[1:-1]), name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            set_run_font(paragraph.add_run(token), color=BLUE, underline=False)
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]))


def add_markdown_table(doc, rows):
    if len(rows) < 2:
        return
    data = rows[0:1] + rows[2:]
    columns = max(len(row) for row in data)
    if columns == 2:
        widths = [2700, 6660]
    elif columns == 3:
        widths = [1800, 3780, 3780]
    else:
        widths = [CONTENT_DXA // columns] * columns
        widths[-1] += CONTENT_DXA - sum(widths)
    table = doc.add_table(rows=len(data), cols=columns)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_props = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_props.append(header_marker)
    for r_index, row in enumerate(data):
        for c_index in range(columns):
            cell = table.cell(r_index, c_index)
            if r_index == 0:
                shade(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.15
            value = row[c_index] if c_index < len(row) else ""
            run = paragraph.add_run(value)
            set_run_font(run, size=9.3, bold=(r_index == 0), color=NAVY if r_index == 0 else BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table_line(line):
    return [part.strip() for part in line.strip().strip("|").split("|")]


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    paragraph._p.get_or_add_pPr().append(shading)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.left_indent = Pt(8)
    paragraph.paragraph_format.right_indent = Pt(8)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    set_run_font(paragraph.add_run("\n".join(lines)), name="Consolas", size=8.5, color="333333")


def new_numbered_list_id(doc: Document, start: int = 1) -> int:
    """Create a fresh Word numbering instance at the Markdown list's stated start."""
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    numbering = doc.part.numbering_part.element
    abstract_id = None
    for num in numbering.num_lst:
        if num.numId == style_num_id:
            abstract_id = num.abstractNumId.val
            break
    if abstract_id is None:
        raise RuntimeError("Could not resolve the List Number numbering definition")
    num = numbering.add_num(abstract_id)
    num.add_lvlOverride(ilvl=0).add_startOverride(start)
    return num.numId


def apply_numbered_list_id(paragraph, num_id: int):
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id


def set_latest_image_alt_text(doc: Document, description: str):
    shape = doc.inline_shapes[-1]
    shape._inline.docPr.set("descr", description)
    shape._inline.docPr.set("title", description)


def markdown_to_doc(doc: Document, path: Path, *, architecture: Path | None = None, ui_map: Path | None = None):
    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    skip_front = True
    in_code = False
    code_lang = ""
    code_lines = []
    inserted_ui = False
    active_numbered_list_id = None
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        is_numbered_line = bool(re.match(r"^\d+\. ", line))
        if not is_numbered_line:
            active_numbered_list_id = None
        if skip_front:
            if line.startswith("## 1."):
                skip_front = False
            else:
                index += 1
                continue
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                if code_lang == "mermaid" and architecture:
                    doc.add_picture(str(architecture), width=Inches(6.5))
                    set_latest_image_alt_text(
                        doc,
                        "Implemented system architecture from public sources through offline processing, SQLite, semantic retrieval, Wiki pages, question answering, and the Streamlit application.",
                    )
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_after = Pt(8)
                    set_run_font(caption.add_run("Figure 1. Implemented architecture and evidence flow."), size=9, color=MUTED, italic=True)
                else:
                    add_code_block(doc, code_lines)
                in_code = False
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1]):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(parse_table_line(lines[index]))
                index += 1
            add_markdown_table(doc, rows)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            if ui_map and not inserted_ui and line.startswith("## 4."):
                doc.add_picture(str(ui_map), width=Inches(6.5))
                set_latest_image_alt_text(
                    doc,
                    "Suggested evaluation flow through Corpus, Search, Wiki, Chatbot, and Evaluation, followed by opening cited sources for important conclusions.",
                )
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(caption.add_run("Figure 1. Suggested path through the five application sections."), size=9, color=MUTED, italic=True)
                inserted_ui = True
            doc.add_heading(line[3:], level=1)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif re.match(r"^- \[[ xX]\] ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, "☐ " + line[6:])
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            if active_numbered_list_id is None:
                stated_start = int(line.split(".", 1)[0])
                active_numbered_list_id = new_numbered_list_id(doc, stated_start)
            apply_numbered_list_id(p, active_numbered_list_id)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.strip():
            paragraph_lines = [line]
            while index + 1 < len(lines):
                nxt = lines[index + 1].strip()
                if not nxt or nxt.startswith(("#", "- ", "|", "```")) or re.match(r"^\d+\. ", nxt):
                    break
                paragraph_lines.append(nxt)
                index += 1
            p = doc.add_paragraph()
            add_inline(p, " ".join(paragraph_lines))
        index += 1


def build_document(markdown: Path, output: Path, *, technical: bool, architecture: Path | None = None, ui_map: Path | None = None):
    doc = Document()
    running = "Conservation Document Intelligence | Technical Handoff" if technical else "Conservation Document Intelligence | MDC Evaluation Guide"
    style_document(doc, running)
    if technical:
        add_cover(
            doc,
            "Technical implementation and handoff",
            "Conservation Document Intelligence Prototype",
            "Architecture, evidence pipeline, safeguards, deployment, and maintenance",
            "Prepared for researchers and technical maintainers",
        )
    else:
        add_cover(
            doc,
            "MDC evaluation and user guide",
            "Conservation Document Intelligence Prototype",
            "How to explore the application, examine citations, and provide useful feedback",
            "Prepared for Missouri Department of Conservation collaborators",
        )
    markdown_to_doc(doc, markdown, architecture=architecture, ui_map=ui_map)
    doc.core_properties.title = "Conservation Document Intelligence Prototype"
    doc.core_properties.subject = "Technical handoff" if technical else "MDC evaluation and user guide"
    doc.core_properties.author = "Conservation Document Intelligence Project"
    doc.core_properties.keywords = "conservation, document intelligence, research prototype"
    doc.save(output)


def main():
    ASSETS.mkdir(parents=True, exist_ok=True)
    architecture = ASSETS / "implemented_system_architecture.png"
    ui_map = ASSETS / "prototype_user_flow.png"
    build_architecture_diagram(architecture)
    build_ui_diagram(ui_map)
    build_document(TECH_MD, TECH_DOCX, technical=True, architecture=architecture)
    build_document(MDC_MD, MDC_DOCX, technical=False, ui_map=ui_map)
    print(TECH_DOCX)
    print(MDC_DOCX)


if __name__ == "__main__":
    main()
